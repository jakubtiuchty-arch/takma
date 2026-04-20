#!/usr/bin/env python3
"""Generate DOCX article for trans.info — Zebra-focused TCO article."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(26, 26, 26)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# === TITLE ===
p = doc.add_paragraph()
run = p.add_run('Prawdziwy koszt sprzętu AutoID w logistyce — cena zakupu to dopiero początek')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(17, 17, 17)
p.paragraph_format.space_after = Pt(4)

# Meta
p = doc.add_paragraph()
run = p.add_run('Artykuł ekspercki · trans.info · marzec 2026')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(136, 136, 136)
p.paragraph_format.space_after = Pt(12)

# === LEAD ===
p = doc.add_paragraph()
run = p.add_run('Terminal mobilny za 2 500 zł. Tyle widzi dział zakupów na fakturze. Ale ile ten terminal kosztuje przez pięć lat codziennej pracy w magazynie? Baterie, naprawy, przestoje, wymiana po trzech latach zamiast po siedmiu. Prawdziwy koszt sprzętu AutoID to zakup plus eksploatacja plus serwis plus utracona produktywność. Kto tego nie liczy, płaci dwa razy.')
run.font.size = Pt(11.5)
run.font.color.rgb = RGBColor(51, 51, 51)
p.paragraph_format.space_after = Pt(14)

# === SECTION 1 ===
doc.add_heading('Czym jest TCO sprzętu AutoID', level=2)

doc.add_paragraph('TCO (Total Cost of Ownership) to suma kosztów od zakupu do wycofania urządzenia. W sprzęcie AutoID składa się z pięciu elementów:')

items1 = [
    ('Zakup', 'urządzenie, akcesoria startowe (stacja dokująca, zapasowa bateria), licencje.'),
    ('Wdrożenie', 'konfiguracja, integracja z WMS lub ERP, szkolenie operatorów.'),
    ('Eksploatacja', 'wymiana baterii, głowic drukujących, zakup etykiet i ribbonów.'),
    ('Serwis', 'naprawy, przeglądy, czas reakcji, urządzenia zastępcze.'),
    ('Wymiana', 'żywotność 3-7 lat, koszt migracji na nowy model.'),
]
for i, (bold, rest) in enumerate(items1, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(bold)
    run.bold = True
    p.add_run(f' — {rest}')

doc.add_paragraph('W logistyce TCO jest krytyczne — sprzęt pracuje w trudnych warunkach (kurz, upadki, mróz), a każda godzina przestoju to realna strata: niewydane zamówienia, opóźnione dostawy, stojący pracownik.')

# === SECTION 2 ===
doc.add_heading('Terminale mobilne — droższy nie znaczy droższy', level=2)

doc.add_paragraph('Scenariusz z wielu magazynów: firma kupuje najtańszy terminal, bo „spełnia wymagania". Po roku bateria trzyma trzy godziny. Po dwóch — pęknięty ekran. Po trzech — koniec wsparcia producenta. Koszt godziny przestoju stanowiska to 80-150 zł (pracownik + utracona przepustowość + kary za opóźnienia).')

doc.add_paragraph('Co decyduje o pięcioletnim koszcie posiadania terminala:')

bullets2 = [
    ('Klasa IP', 'IP67/IP68 (pyłoszczelność + zanurzenie) = statystycznie mniej napraw niż IP54/65.'),
    ('Odporność na upadki', 'MIL-STD-810G/H, 1,5-3,0 m na beton. Każdy przeżyty upadek to zaoszczędzone 400-1 200 zł.'),
    ('Bateria hot-swap', 'wymiana bez wyłączania eliminuje przestoje na ładowanie.'),
    ('Wsparcie producenta', 'Zebra LifeGuard daje aktualizacje do 10 lat; budżetowe marki kończą po 3 latach.'),
]
for bold, rest in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold)
    run.bold = True
    p.add_run(f' — {rest}')

p = doc.add_paragraph()
p.add_run('Porównanie na urządzeniach Zebra: TC22 za 2 417 zł (IP68, 6", Wi-Fi 6E) po pięciu latach generuje TCO 5 200-6 500 zł — wliczając wymianę baterii i naprawy. MC3400 za 4 561 zł (gun z klawiaturą, MIL-STD-810H) — TCO 6 000-7 000 zł, ale z minimalną liczbą serwisów i siedmioletnią żywotnością. Cena zakupu podwójna — TCO rocznie porównywalne. Porównując ')
run = p.add_run('terminale mobilne do magazynu')
run.font.color.rgb = RGBColor(26, 86, 219)
run.underline = True
p.add_run(' [https://takma.com.pl/terminale-mobilne] warto zestawić nie cenę z faktury, a pięcioletni koszt posiadania.')

# === SECTION 3 ===
doc.add_heading('Drukarki etykiet — koszt ukryty w każdej przesyłce', level=2)

doc.add_paragraph('Drukarka etykiet to ciągły koszt operacyjny. Biurkowa Zebra ZD421t (od 1 649 zł) w trybie termicznym (DT) nie wymaga ribbonu, ale etykiety blakną — do etykiet wysyłkowych. Przemysłowa Zebra ZT231 (od 2 550 zł) w termotransferze (TT) wymaga ribbonu, ale etykiety trwają latami — do regałów i inwentaryzacji. Wybór metody druku wpływa na TCO bardziej niż cena drukarki.')

doc.add_paragraph('Największy ukryty koszt to głowica drukująca — wymiana co 50-150 km druku, 300-800 zł (głowica do Zebra ZT411: ok. 550 zł). Tanie materiały skracają żywotność głowicy — pozorna oszczędność 20% oznacza dwukrotnie częstszą wymianę. Przy 1 000 etykiet dziennie to tysiące złotych rocznie.')

p = doc.add_paragraph()
p.add_run('Koszt jednej etykiety = amortyzacja drukarki + materiał + zużycie głowicy. Mało kto to liczy. Przy wyborze ')
run = p.add_run('drukarki etykiet dla logistyki')
run.font.color.rgb = RGBColor(26, 86, 219)
run.underline = True
p.add_run(' [https://takma.com.pl/drukarki-etykiet] kluczowy jest łączny koszt eksploatacji, nie cena katalogowa.')

# === SECTION 4 ===
doc.add_heading('Serwis — czynnik, który przesądza o TCO', level=2)

doc.add_paragraph('Terminal w naprawie przez pięć dni roboczych to 3 000-6 000 zł utraconych (przy 80-150 zł/h przestoju stanowiska). Serwis z reakcją next business day i urządzeniem zastępczym ogranicza ten koszt do minimum. Pakiet serwisowy za 300-600 zł rocznie jest wielokrotnie tańszy niż koszt jednego dłuższego przestoju.')

doc.add_paragraph('Cztery rzeczy do sprawdzenia przy wyborze partnera serwisowego:')

items4 = [
    ('Czas reakcji (SLA)', 'czy gwarantuje next business day, czy „do 14 dni roboczych"?'),
    ('Autoryzacja producenta', 'dostęp do oryginalnych części, firmware i narzędzi diagnostycznych.'),
    ('Urządzenia zastępcze', 'czy na czas naprawy otrzymasz działający terminal?'),
    ('Serwis w Polsce', 'lokalna naprawa trwa dni; wysyłka za granicę — tygodnie.'),
]
for i, (bold, rest) in enumerate(items4, 1):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(bold)
    run.bold = True
    p.add_run(f' — {rest}')

p = doc.add_paragraph()
p.add_run('Zebra Technologies to najczęściej spotykany producent sprzętu AutoID w polskiej logistyce. Dostęp do ')
run = p.add_run('autoryzowanego serwisu urządzeń Zebra')
run.font.color.rgb = RGBColor(26, 86, 219)
run.underline = True
p.add_run(' [https://serwis-zebry.pl] z gwarantowanym SLA i oryginalnymi częściami bezpośrednio obniża TCO floty.')

# === SUMMARY ===
doc.add_heading('Podsumowanie', level=2)

doc.add_paragraph('Reguła kciuka: TCO sprzętu AutoID to 2-4 razy cena zakupu. Przed zakupem warto zadać trzy pytania: ile kosztuje eksploatacja przez pięć lat? Jak długo producent wspiera urządzenie? Jaki jest czas reakcji serwisu?')

p = doc.add_paragraph()
run = p.add_run('Kto liczy TCO przed zakupem — oszczędza. Kto nie — płaci dwa razy.')
run.italic = True
run.font.color.rgb = RGBColor(68, 68, 68)

# === LINKS REFERENCE ===
doc.add_paragraph('')  # spacer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run('LINKI DOFOLLOW W ARTYKULE (3 szt.):')
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(100, 100, 100)

links = [
    '1. „terminale mobilne do magazynu" → https://takma.com.pl/terminale-mobilne (sekcja 2)',
    '2. „drukarki etykiet dla logistyki" → https://takma.com.pl/drukarki-etykiet (sekcja 3)',
    '3. „autoryzowanego serwisu urządzeń Zebra" → https://serwis-zebry.pl (sekcja 4)',
]
for link in links:
    p = doc.add_paragraph()
    run = p.add_run(link)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.space_after = Pt(0)

out = '/Users/jakubtiuchty/takma/public/downloads/ART-tco-sprzetu-autoid-trans-info.docx'
doc.save(out)
print(f'DOCX saved: {out}')
