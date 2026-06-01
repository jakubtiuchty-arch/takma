#!/usr/bin/env python3
"""Convert article markdown to .docx with proper hyperlinks and tables."""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

SRC = "/Users/jakubtiuchty/takma/drafts/artykul-agrofakty-paszporty-roslin.md"
OUT = "/Users/jakubtiuchty/takma/drafts/artykul-agrofakty-paszporty-roslin.docx"


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_RE = re.compile(r"(\[([^\]]+)\]\(([^)]+)\))|(\*\*([^*]+)\*\*)")


def add_inline_runs(paragraph, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1):
            add_hyperlink(paragraph, m.group(3), m.group(2))
        elif m.group(4):
            run = paragraph.add_run(m.group(5))
            run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table_row(line):
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts


def is_table_separator(line):
    s = line.strip().strip("|").replace(" ", "")
    return bool(s) and all(c in "-:" for c in s)


def main():
    with open(SRC, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph().add_run("─" * 40)
            i += 1
            continue

        # Markdown table: |col|col| then |---|---|
        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h.replace("**", ""))
                run.bold = True
            for r_idx, row in enumerate(rows):
                for j, val in enumerate(row):
                    if j < len(table.rows[r_idx + 1].cells):
                        cell = table.rows[r_idx + 1].cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_runs(p, val)
            doc.add_paragraph()
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, line)
        i += 1

    doc.save(OUT)
    print(f"WROTE {OUT}")


if __name__ == "__main__":
    main()
